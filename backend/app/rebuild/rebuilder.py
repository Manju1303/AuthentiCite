import os
from docx import Document
from typing import List, Dict, Any
from backend.app.rebuild.styles.original import apply_original_layout
from backend.app.rebuild.styles.ieee import apply_ieee_layout
from backend.app.rebuild.styles.springer import apply_springer_layout

def rebuild_document(
    sections: List[Dict[str, Any]],
    references: List[Dict[str, Any]],
    layout_map: Dict[str, Any],
    output_path: str,
    journal_format: str = "original",
    media_dir: str = "uploads/media"
) -> str:
    """
    Creates a new docx document, applies the requested journal style layout,
    inserts all sections and references, and saves the file.
    """
    doc = Document()
    
    # Remove the default paragraph that python-docx adds to new documents
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
        
    journal_format = journal_format.lower().strip()
    
    if journal_format == "ieee":
        apply_ieee_layout(doc, sections, references, layout_map, media_dir)
    elif journal_format == "springer":
        apply_springer_layout(doc, sections, references, layout_map, media_dir)
    else:
        apply_original_layout(doc, sections, references, layout_map, media_dir)
        
    # Ensure directory exists and save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
