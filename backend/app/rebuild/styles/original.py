import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any

def apply_original_layout(
    doc: Document, 
    sections: List[Dict[str, Any]], 
    references: List[Dict[str, Any]], 
    layout_map: Dict[str, Any],
    media_dir: str = "uploads/media"
):
    """
    Rebuilds the document in docx preserving the original layout styling
    (fonts, bold/italic, size, alignment, margins) paragraph by paragraph.
    """
    # 1. Apply page margins
    if "margins" in layout_map and doc.sections:
        m = layout_map["margins"]
        for section in doc.sections:
            section.top_margin = Inches(m.get("top", 1.0))
            section.bottom_margin = Inches(m.get("bottom", 1.0))
            section.left_margin = Inches(m.get("left", 1.0))
            section.right_margin = Inches(m.get("right", 1.0))
            
    # 2. Add elements
    for block in sections:
        meta = block["layout_metadata"]
        b_type = meta.get("type", "paragraph")
        
        # Check alignment
        align = WD_ALIGN_PARAGRAPH.LEFT
        align_str = meta.get("alignment", "LEFT")
        if align_str == "CENTER":
            align = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "RIGHT":
            align = WD_ALIGN_PARAGRAPH.RIGHT
        elif align_str == "JUSTIFY":
            align = WD_ALIGN_PARAGRAPH.JUSTIFY

        if b_type == "paragraph":
            p = doc.add_paragraph()
            p.alignment = align
            
            # Use spacing from metadata
            p.paragraph_format.space_before = Pt(meta.get("spacing_before", 0))
            p.paragraph_format.space_after = Pt(meta.get("spacing_after", 6))
            p.paragraph_format.line_spacing = meta.get("line_spacing", 1.15)
            
            # Reconstruct runs
            runs = meta.get("runs", [])
            text_to_write = block.get("rewritten_text") or block.get("original_text")
            
            # If we rewrote the text, we might have a single string instead of split runs.
            # Let's map the rewritten text to the original runs format, 
            # or write it as a single styled run if runs are missing or mismatched.
            if len(runs) <= 1 or not block.get("rewritten_text"):
                r = p.add_run(text_to_write)
                if runs:
                    r.font.name = runs[0].get("font_name", "Calibri")
                    r.font.size = Pt(runs[0].get("font_size", 11))
                    r.bold = runs[0].get("bold", False)
                    r.italic = runs[0].get("italic", False)
                    color_hex = runs[0].get("color")
                    if color_hex:
                        color_hex = color_hex.lstrip('#')
                        r.font.color.rgb = RGBColor.from_string(color_hex)
            else:
                # Apply style attributes of the first run as fallback, but write full text
                # to prevent cutting off text if length changed during rewriting.
                r = p.add_run(text_to_write)
                r.font.name = runs[0].get("font_name", "Calibri")
                r.font.size = Pt(runs[0].get("font_size", 11))
                r.bold = runs[0].get("bold", False)
                r.italic = runs[0].get("italic", False)
                color_hex = runs[0].get("color")
                if color_hex:
                    color_hex = color_hex.lstrip('#')
                    r.font.color.rgb = RGBColor.from_string(color_hex)
                    
        elif b_type == "image":
            image_name = meta.get("image_name")
            image_path = os.path.join(media_dir, image_name)
            if os.path.exists(image_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(image_path, width=Inches(5.0))
                
        elif b_type == "table":
            text = block.get("original_text") # Tables are kept intact
            # Reconstruct table from markdown representation
            lines = text.strip().split("\n")
            if len(lines) >= 2:
                # Parse markdown back to cell grid
                rows_cells = []
                for line in lines:
                    if "---" in line and "|" in line:
                        continue # skip separator row
                    cells = [c.strip() for c in line.split("|")]
                    if cells:
                        rows_cells.append(cells)
                
                if rows_cells:
                    table = doc.add_table(rows=len(rows_cells), cols=len(rows_cells[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows_cells):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(table.columns):
                                table.cell(r_idx, c_idx).text = val
            p = doc.add_paragraph() # spacing paragraph after table
            p.paragraph_format.space_before = Pt(6)

    # 3. Add References
    if references:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run("References")
        r.font.name = "Calibri"
        r.font.size = Pt(16)
        r.bold = True
        
        for ref in references:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(4)
            p_ref.paragraph_format.left_indent = Inches(0.25)
            r_ref = p_ref.add_run(ref["raw_reference"])
            r_ref.font.name = "Calibri"
            r_ref.font.size = Pt(11)
