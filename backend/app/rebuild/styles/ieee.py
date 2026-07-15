import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any

def apply_ieee_layout(
    doc: Document, 
    sections: List[Dict[str, Any]], 
    references: List[Dict[str, Any]], 
    layout_map: Dict[str, Any],
    media_dir: str = "uploads/media"
):
    """
    Applies the official IEEE journal/conference paper format to the rebuilt document:
    - Times New Roman font throughout
    - Single-column Title, Authors, and Abstract/Keywords
    - Double-column body layout
    - Roman numeral centered headings (Level 1)
    - Italic left-aligned subheadings (Level 2)
    - 10pt justified body text with 0.15 in first line indent
    - 8pt references bibliography list
    """
    # 1. Setup Page Margins on the first section (Title & Abstract)
    first_section = doc.sections[0]
    first_section.top_margin = Inches(0.75)
    first_section.bottom_margin = Inches(1.0)
    first_section.left_margin = Inches(0.625)
    first_section.right_margin = Inches(0.625)
    
    # Track which blocks go into single-column (Title, Authors, Abstract, Keywords)
    title_block = None
    abstract_block = None
    keywords_block = None
    body_blocks = []
    
    for block in sections:
        sec_name = block["section_name"]
        text = block.get("rewritten_text") or block.get("original_text")
        meta = block["layout_metadata"]
        
        # Classify sections for layout routing
        if sec_name == "Title/Abstract" and not title_block:
            title_block = block
        elif sec_name == "Abstract":
            abstract_block = block
        elif sec_name == "Keywords":
            keywords_block = block
        else:
            body_blocks.append(block)

    # Helper function to add structured Times New Roman runs
    def add_styled_paragraph(doc_obj, text_content, font_size_pt, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, first_line_indent=0):
        p = doc_obj.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.0  # IEEE is single-spaced
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Inches(first_line_indent)
            
        r = p.add_run(text_content)
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size_pt)
        r.bold = bold
        r.italic = italic
        return p

    # --- 1. WRITE SINGLE-COLUMN HEADER ---
    
    # A. Title (24pt)
    title_text = title_block.get("rewritten_text") or title_block.get("original_text") if title_block else "Research Paper Title"
    # Filter title if it contains author details or split them
    lines = title_text.split("\n")
    paper_title = lines[0]
    author_info = "\n".join(lines[1:]) if len(lines) > 1 else "Author Name\nAuthor Affiliation"
    
    add_styled_paragraph(doc, paper_title, 24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    
    # B. Authors (11pt)
    add_styled_paragraph(doc, author_info, 11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)
    
    # C. Abstract & Keywords (9pt bold italic)
    if abstract_block:
        abs_text = abstract_block.get("rewritten_text") or abstract_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        
        r_lbl = p.add_run("Abstract—")
        r_lbl.font.name = "Times New Roman"
        r_lbl.font.size = Pt(9)
        r_lbl.bold = True
        r_lbl.italic = True
        
        r_val = p.add_run(abs_text)
        r_val.font.name = "Times New Roman"
        r_val.font.size = Pt(9)
        r_val.italic = True
        
    if keywords_block:
        kw_text = keywords_block.get("rewritten_text") or keywords_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(12)
        
        r_lbl = p.add_run("Keywords—")
        r_lbl.font.name = "Times New Roman"
        r_lbl.font.size = Pt(9)
        r_lbl.bold = True
        r_lbl.italic = True
        
        r_val = p.add_run(kw_text)
        r_val.font.name = "Times New Roman"
        r_val.font.size = Pt(9)
        r_val.italic = True

    # --- 2. START DOUBLE-COLUMN BODY SECTION ---
    body_section = doc.add_section()
    body_section.top_margin = Inches(0.75)
    body_section.bottom_margin = Inches(1.0)
    body_section.left_margin = Inches(0.625)
    body_section.right_margin = Inches(0.625)
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Set columns via XML properties to bypass python-docx API limitation
    sectPr = body_section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    else:
        cols = cols[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360') # 360 twips = 0.25 inches

    
    # Regex to recognize structural headings in the body
    # e.g., "1. Introduction" or "I. Introduction"
    level1_regex = r"^([IXV\d]+\.\s+)"
    level2_regex = r"^([A-Z]\.\s+)"
    
    roman_numeral_map = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    l1_counter = 0

    for block in body_blocks:
        meta = block["layout_metadata"]
        b_type = meta.get("type", "paragraph")
        text = block.get("rewritten_text") or block.get("original_text")
        
        if b_type == "paragraph":
            # Determine if this paragraph is actually a heading
            is_heading_1 = False
            is_heading_2 = False
            
            # Simple heuristic: bold and short, or matching heading pattern
            clean_text = text.strip()
            word_count = len(clean_text.split())
            
            if re.match(level1_regex, clean_text) or (word_count < 6 and block["section_name"] != "Title/Abstract" and meta.get("style", "").startswith("Heading")):
                is_heading_1 = True
            elif re.match(level2_regex, clean_text):
                is_heading_2 = True
                
            if is_heading_1:
                l1_counter += 1
                # Format to ALL CAPS and prepend Roman numeral if not present
                header_text = clean_text.upper()
                if not re.match(r"^[IXV]+\.", header_text):
                    numeral = roman_numeral_map[l1_counter - 1] if l1_counter <= len(roman_numeral_map) else f"{l1_counter}."
                    header_text = f"{numeral}. {header_text}"
                
                add_styled_paragraph(doc, header_text, 10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
                
            elif is_heading_2:
                # Subheading italic left-aligned
                add_styled_paragraph(doc, clean_text, 10, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4)
                
            else:
                # Normal paragraph text (10pt, justified, 0.15" first line indent)
                add_styled_paragraph(doc, text, 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0, first_line_indent=0.15)
                
        elif b_type == "image":
            image_name = meta.get("image_name")
            image_path = os.path.join(media_dir, image_name)
            if os.path.exists(image_path):
                # Double column layout images are usually centered in column width
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(image_path, width=Inches(3.2)) # fits standard IEEE column width (3.2 inches)
                
                # Bottom caption
                add_styled_paragraph(doc, f"Fig. {element_index}. Fig Caption", 8, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
                
        elif b_type == "table":
            # Reconstruct table centered
            lines = text.strip().split("\n")
            if len(lines) >= 2:
                rows_cells = []
                for line in lines:
                    if "---" in line and "|" in line:
                        continue
                    cells = [c.strip() for c in line.split("|")]
                    if cells:
                        rows_cells.append(cells)
                
                if rows_cells:
                    # IEEE captions are ABOVE tables in 8pt Small Caps (or capitalized)
                    add_styled_paragraph(doc, f"TABLE {element_index}. TABLE CAPTION", 8, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
                    
                    table = doc.add_table(rows=len(rows_cells), cols=len(rows_cells[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows_cells):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = val
                                # Style cell paragraphs in 8pt Times New Roman
                                for cp in cell.paragraphs:
                                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for cr in cp.runs:
                                        cr.font.name = "Times New Roman"
                                        cr.font.size = Pt(8)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)

    # --- 3. WRITE REFERENCES BIBLIOGRAPHY SECTION (8pt Times New Roman) ---
    if references:
        # Roman numeral heading for References
        l1_counter += 1
        numeral = roman_numeral_map[l1_counter - 1] if l1_counter <= len(roman_numeral_map) else f"{l1_counter}."
        add_styled_paragraph(doc, f"{numeral}. REFERENCES", 10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
        
        for idx, ref in enumerate(references):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.first_line_indent = Inches(-0.15) # hanging indent
            p.paragraph_format.space_after = Pt(2)
            
            # Format numeric key [1]
            cit_key = ref.get("citation_key") or str(idx + 1)
            r_key = p.add_run(f"[{cit_key}] ")
            r_key.font.name = "Times New Roman"
            r_key.font.size = Pt(8)
            
            # Format reference text
            raw_ref = ref["raw_reference"]
            # Clean off leading brackets if present to avoid duplicate
            raw_ref_clean = re.sub(r"^\[\d+\]\s*", "", raw_ref)
            
            r_ref = p.add_run(raw_ref_clean)
            r_ref.font.name = "Times New Roman"
            r_ref.font.size = Pt(8)
