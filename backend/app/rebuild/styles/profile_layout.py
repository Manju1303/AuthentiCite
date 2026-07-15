# backend/app/rebuild/styles/profile_layout.py

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any

def apply_profile_layout(
    doc: Document, 
    sections: List[Dict[str, Any]], 
    references: List[Dict[str, Any]], 
    layout_map: Dict[str, Any],
    profile: Dict[str, Any],
    media_dir: str = "uploads/media"
):
    """
    Applies the layout of the selected journal profile to the rebuilt document.
    """
    # 1. Setup Page Margins on the first section
    margins = profile.get("margins_in", {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0})
    first_section = doc.sections[0]
    first_section.top_margin = Inches(margins.get("top", 1.0))
    first_section.bottom_margin = Inches(margins.get("bottom", 1.0))
    first_section.left_margin = Inches(margins.get("left", 1.0))
    first_section.right_margin = Inches(margins.get("right", 1.0))

    title_block = None
    abstract_block = None
    keywords_block = None
    body_blocks = []
    
    for block in sections:
        sec_name = block["section_name"]
        if sec_name == "Title/Abstract" and not title_block:
            title_block = block
        elif sec_name == "Abstract":
            abstract_block = block
        elif sec_name == "Keywords":
            keywords_block = block
        else:
            body_blocks.append(block)

    # Alignment lookup map
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    # Helper function to add styled paragraphs with the profile font
    def add_styled_paragraph(doc_obj, text_content, font_size_pt, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, first_line_indent=0):
        p = doc_obj.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = profile.get("body_line_spacing", 1.15)
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Inches(first_line_indent)
            
        r = p.add_run(text_content)
        r.font.name = profile.get("font_name", "Times New Roman")
        r.font.size = Pt(font_size_pt)
        r.bold = bold
        r.italic = italic
        return p

    # --- 1. WRITE SINGLE-COLUMN HEADER ---
    
    # A. Title
    title_text = title_block.get("rewritten_text") or title_block.get("original_text") if title_block else "Research Paper Title"
    lines = title_text.split("\n")
    paper_title = lines[0]
    author_info = "\n".join(lines[1:]) if len(lines) > 1 else "Author Name\nAuthor Affiliation"
    
    title_align = align_map.get(profile.get("title_align", "CENTER"), WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(
        doc, 
        paper_title, 
        profile.get("title_size", 18), 
        bold=profile.get("title_bold", True), 
        align=title_align, 
        space_before=12, 
        space_after=12
    )
    
    # B. Authors
    author_align = align_map.get(profile.get("author_align", "CENTER"), WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(
        doc, 
        author_info, 
        profile.get("author_size", 11), 
        align=author_align, 
        space_before=6, 
        space_after=18
    )
    
    # C. Abstract & Keywords
    if abstract_block:
        abs_text = abstract_block.get("rewritten_text") or abstract_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        abs_indent = profile.get("abstract_indent_in", 0.0)
        if abs_indent > 0:
            p.paragraph_format.left_indent = Inches(abs_indent)
            p.paragraph_format.right_indent = Inches(abs_indent)
        p.paragraph_format.space_after = Pt(6)
        
        lbl_text = profile.get("abstract_label", "")
        if lbl_text:
            r_lbl = p.add_run(lbl_text)
            r_lbl.font.name = profile.get("font_name", "Times New Roman")
            r_lbl.font.size = Pt(profile.get("abstract_size", 10))
            r_lbl.bold = True
            
        r_val = p.add_run(abs_text)
        r_val.font.name = profile.get("font_name", "Times New Roman")
        r_val.font.size = Pt(profile.get("abstract_size", 10))
        r_val.italic = profile.get("abstract_italic", False)
        
    if keywords_block:
        kw_text = keywords_block.get("rewritten_text") or keywords_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        abs_indent = profile.get("abstract_indent_in", 0.0)
        if abs_indent > 0:
            p.paragraph_format.left_indent = Inches(abs_indent)
            p.paragraph_format.right_indent = Inches(abs_indent)
        p.paragraph_format.space_after = Pt(12)
        
        lbl_text = profile.get("keywords_label", "")
        if lbl_text:
            r_lbl = p.add_run(lbl_text)
            r_lbl.font.name = profile.get("font_name", "Times New Roman")
            r_lbl.font.size = Pt(profile.get("abstract_size", 10))
            r_lbl.bold = True
            
        r_val = p.add_run(kw_text)
        r_val.font.name = profile.get("font_name", "Times New Roman")
        r_val.font.size = Pt(profile.get("abstract_size", 10))
        r_val.italic = profile.get("abstract_italic", False)

    # --- 2. START BODY SECTION ---
    if profile.get("body_columns_num", 1) > 1:
        body_section = doc.add_section()
        body_section.top_margin = Inches(margins.get("top", 1.0))
        body_section.bottom_margin = Inches(margins.get("bottom", 1.0))
        body_section.left_margin = Inches(margins.get("left", 1.0))
        body_section.right_margin = Inches(margins.get("right", 1.0))
        
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Set columns via XML properties to bypass python-docx API limitation
        sectPr = body_section._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]
        
        cols.set(qn('w:num'), str(profile.get("body_columns_num", 1)))
        gap_twips = int(profile.get("body_columns_gap_in", 0.25) * 1440)
        cols.set(qn('w:space'), str(gap_twips))


    roman_numeral_map = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX"]
    l1_counter = 0
    l2_counter = 0
    fig_counter = 0
    tbl_counter = 0

    def strip_numbering(t):
        return re.sub(r"^([IXV\d]+\.\d*\s*|[A-Z]\.\s*|[\d\.]+\s*)", "", t.strip())

    for block in body_blocks:
        meta = block["layout_metadata"]
        b_type = meta.get("type", "paragraph")
        text = block.get("rewritten_text") or block.get("original_text")
        
        if b_type == "paragraph":
            clean_text = text.strip()
            word_count = len(clean_text.split())
            style_name = meta.get("style", "")
            
            # Heading heuristics
            is_heading_1 = False
            is_heading_2 = False
            
            if (
                style_name.startswith("Heading 1") or 
                (word_count < 8 and block.get("section_name") != "Title/Abstract" and (
                    re.match(r"^([IXV]+|\d+)\.\s*", clean_text) or 
                    re.match(r"^[A-Z]\.\s*", clean_text) or 
                    style_name == "Heading 1"
                ))
            ):
                is_heading_1 = True
            elif (
                style_name.startswith("Heading 2") or 
                (word_count < 8 and (
                    re.match(r"^(\d+\.\d+|[A-Z]\.[0-9]+)\s*", clean_text) or 
                    style_name == "Heading 2"
                ))
            ):
                is_heading_2 = True
                
            if is_heading_1:
                l1_counter += 1
                l2_counter = 0
                header_text = strip_numbering(clean_text)
                
                num_style = profile.get("heading1_numbering", "none")
                if num_style == "roman":
                    num = roman_numeral_map[l1_counter - 1] if l1_counter <= len(roman_numeral_map) else str(l1_counter)
                    header_text = f"{num}. {header_text}"
                elif num_style == "decimal":
                    header_text = f"{l1_counter}. {header_text}"
                elif num_style == "decimal_no_dot":
                    header_text = f"{l1_counter} {header_text}"
                    
                casing = profile.get("heading1_casing", "none")
                if casing == "upper":
                    header_text = header_text.upper()
                elif casing == "title":
                    header_text = header_text.title()
                
                add_styled_paragraph(
                    doc, 
                    header_text, 
                    profile.get("heading1_size", 12),
                    bold=profile.get("heading1_bold", True),
                    italic=profile.get("heading1_italic", False),
                    align=align_map.get(profile.get("heading1_align", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT),
                    space_before=12,
                    space_after=6
                )
                
            elif is_heading_2:
                l2_counter += 1
                header_text = strip_numbering(clean_text)
                
                num_style = profile.get("heading2_numbering", "none")
                if num_style == "decimal_dot":
                    header_text = f"{l1_counter}.{l2_counter} {header_text}"
                elif num_style == "alpha":
                    num = chr(64 + l2_counter) if l2_counter <= 26 else str(l2_counter)
                    header_text = f"{num}. {header_text}"
                    
                casing = profile.get("heading2_casing", "none")
                if casing == "upper":
                    header_text = header_text.upper()
                elif casing == "title":
                    header_text = header_text.title()
                    
                add_styled_paragraph(
                    doc, 
                    header_text, 
                    profile.get("heading2_size", 10),
                    bold=profile.get("heading2_bold", True),
                    italic=profile.get("heading2_italic", False),
                    align=align_map.get(profile.get("heading2_align", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT),
                    space_before=6,
                    space_after=4
                )
                
            else:
                # Normal paragraph text
                add_styled_paragraph(
                    doc, 
                    text, 
                    profile.get("body_size", 10),
                    align=align_map.get(profile.get("body_align", "JUSTIFY"), WD_ALIGN_PARAGRAPH.JUSTIFY),
                    space_before=0,
                    space_after=profile.get("body_space_after_pt", 0.0),
                    first_line_indent=profile.get("body_first_line_indent_in", 0.15)
                )
                
        elif b_type == "image":
            image_name = meta.get("image_name")
            image_path = os.path.join(media_dir, image_name)
            if os.path.exists(image_path):
                fig_counter += 1
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                width_in = 3.2 if profile.get("body_columns_num", 1) > 1 else 5.0
                p.add_run().add_picture(image_path, width=Inches(width_in))
                
                # Bottom caption
                add_styled_paragraph(
                    doc, 
                    f"Fig. {fig_counter}. Figure Caption", 
                    profile.get("body_size", 10) - 1, 
                    align=WD_ALIGN_PARAGRAPH.CENTER, 
                    space_before=4, 
                    space_after=8
                )
                
        elif b_type == "table":
            lines = text.strip().split("\n")
            if len(lines) >= 2:
                tbl_counter += 1
                rows_cells = []
                for line in lines:
                    if "---" in line and "|" in line:
                        continue
                    cells = [c.strip() for c in line.split("|")]
                    if cells:
                        rows_cells.append(cells)
                
                if rows_cells:
                    # Caption is above tables
                    add_styled_paragraph(
                        doc, 
                        f"Table {tbl_counter}. Table Caption", 
                        profile.get("body_size", 10) - 1, 
                        align=WD_ALIGN_PARAGRAPH.CENTER, 
                        space_before=8, 
                        space_after=4
                    )
                    
                    table = doc.add_table(rows=len(rows_cells), cols=len(rows_cells[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows_cells):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = val
                                for cp in cell.paragraphs:
                                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for cr in cp.runs:
                                        cr.font.name = profile.get("font_name", "Times New Roman")
                                        cr.font.size = Pt(profile.get("body_size", 10) - 1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)

    # --- 3. WRITE REFERENCES BIBLIOGRAPHY SECTION ---
    if references:
        add_styled_paragraph(
            doc, 
            "References" if profile.get("name") != "Journal of the ACM" else "REFERENCES", 
            profile.get("heading1_size", 12), 
            bold=True, 
            align=align_map.get(profile.get("heading1_align", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT),
            space_before=18,
            space_after=8
        )
        
        for idx, ref in enumerate(references):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Hanging indent
            hanging = profile.get("references_hanging_indent_in", 0.0)
            if hanging > 0:
                p.paragraph_format.left_indent = Inches(hanging)
                p.paragraph_format.first_line_indent = Inches(-hanging)
                
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = profile.get("body_line_spacing", 1.15)
            
            # Format number/citation key
            num_style = profile.get("references_numbering", "none")
            cit_key = ref.get("citation_key") or str(idx + 1)
            
            if num_style == "brackets":
                r_key = p.add_run(f"[{cit_key}] ")
                r_key.font.name = profile.get("font_name", "Times New Roman")
                r_key.font.size = Pt(profile.get("references_size", 9))
                r_key.bold = True
            elif num_style == "number_dot":
                r_key = p.add_run(f"{cit_key}. ")
                r_key.font.name = profile.get("font_name", "Times New Roman")
                r_key.font.size = Pt(profile.get("references_size", 9))
                r_key.bold = True
                
            raw_ref = ref["raw_reference"]
            if num_style in ["brackets", "number_dot"]:
                raw_ref = re.sub(r"^\[?\d+\]?\s*\.?\s*", "", raw_ref)
                
            r_ref = p.add_run(raw_ref)
            r_ref.font.name = profile.get("font_name", "Times New Roman")
            r_ref.font.size = Pt(profile.get("references_size", 9))
