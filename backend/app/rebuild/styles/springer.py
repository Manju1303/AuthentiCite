import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any

def apply_springer_layout(
    doc: Document, 
    sections: List[Dict[str, Any]], 
    references: List[Dict[str, Any]], 
    layout_map: Dict[str, Any],
    media_dir: str = "uploads/media"
):
    """
    Applies the Springer LNCS (Lecture Notes in Computer Science) layout format:
    - Single-column throughout
    - Times New Roman font
    - 1.2 in page margins
    - Title: 16pt bold centered
    - Author: 10pt centered
    - Abstract: 9pt justified, 0.4 in left/right indent
    - Headings: Bold numbered left-aligned (12pt for Level 1, 10pt for Level 2)
    - Body text: 10pt justified, 0.15 in first line indent, single spacing
    - References: 9pt numbered list
    """
    # 1. Setup Page Margins
    for sec in doc.sections:
        sec.top_margin = Inches(1.2)
        sec.bottom_margin = Inches(1.2)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)
        
    title_block = None
    abstract_block = None
    keywords_block = None
    body_blocks = []
    
    for block in sections:
        sec_name = block["section_name"]
        if sec_name == "Title/Abstract" and not title_block:
            title_block = block
        elif sec_name == "Abstract":
            abstract_block = block
        elif sec_name == "Keywords":
            keywords_block = block
        else:
            body_blocks.append(block)

    def add_styled_paragraph(doc_obj, text_content, font_size_pt, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, first_line_indent=0):
        p = doc_obj.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.0
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Inches(first_line_indent)
            
        r = p.add_run(text_content)
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size_pt)
        r.bold = bold
        r.italic = italic
        return p

    # --- 1. TITLE AND AUTHOR ---
    title_text = title_block.get("rewritten_text") or title_block.get("original_text") if title_block else "Research Paper Title"
    lines = title_text.split("\n")
    paper_title = lines[0]
    author_info = "\n".join(lines[1:]) if len(lines) > 1 else "Author Name\nAuthor Affiliation"
    
    add_styled_paragraph(doc, paper_title, 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    add_styled_paragraph(doc, author_info, 10, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)
    
    # --- 2. ABSTRACT & KEYWORDS ---
    if abstract_block:
        abs_text = abstract_block.get("rewritten_text") or abstract_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(6)
        
        r_lbl = p.add_run("Abstract. ")
        r_lbl.font.name = "Times New Roman"
        r_lbl.font.size = Pt(9)
        r_lbl.bold = True
        
        r_val = p.add_run(abs_text)
        r_val.font.name = "Times New Roman"
        r_val.font.size = Pt(9)
        
    if keywords_block:
        kw_text = keywords_block.get("rewritten_text") or keywords_block.get("original_text")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(18)
        
        r_lbl = p.add_run("Keywords: ")
        r_lbl.font.name = "Times New Roman"
        r_lbl.font.size = Pt(9)
        r_lbl.bold = True
        
        r_val = p.add_run(kw_text)
        r_val.font.name = "Times New Roman"
        r_val.font.size = Pt(9)

    # --- 3. BODY SECTIONS ---
    l1_counter = 0
    l2_counter = 0
    
    for block in body_blocks:
        meta = block["layout_metadata"]
        b_type = meta.get("type", "paragraph")
        text = block.get("rewritten_text") or block.get("original_text")
        
        if b_type == "paragraph":
            clean_text = text.strip()
            word_count = len(clean_text.split())
            
            # Heading identification
            is_heading_1 = False
            is_heading_2 = False
            
            if re.match(r"^(\d+\.\s+)", clean_text) or (word_count < 6 and block["section_name"] != "Title/Abstract" and meta.get("style", "").startswith("Heading 1")):
                is_heading_1 = True
            elif re.match(r"^(\d+\.\d+\s+)", clean_text) or (word_count < 6 and meta.get("style", "").startswith("Heading 2")):
                is_heading_2 = True
                
            if is_heading_1:
                l1_counter += 1
                l2_counter = 0 # reset subheading counter
                header_text = clean_text
                # Enforce numbering e.g. "1 Introduction" (without trailing dot in Springer LNCS)
                header_text = re.sub(r"^[\d\.]+\s*", "", header_text) # remove existing number
                header_text = f"{l1_counter} {header_text}"
                
                add_styled_paragraph(doc, header_text, 12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
                
            elif is_heading_2:
                l2_counter += 1
                header_text = clean_text
                header_text = re.sub(r"^[\d\.]+\s*", "", header_text)
                header_text = f"{l1_counter}.{l2_counter} {header_text}"
                
                add_styled_paragraph(doc, header_text, 10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4)
                
            else:
                # Normal paragraph text (10pt, justified, first line indent = 0.15 in)
                add_styled_paragraph(doc, text, 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0, first_line_indent=0.15)
                
        elif b_type == "image":
            image_name = meta.get("image_name")
            image_path = os.path.join(media_dir, image_name)
            if os.path.exists(image_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(image_path, width=Inches(4.5))
                
                # Figure caption below image
                add_styled_paragraph(doc, f"Fig. {element_index}. Fig Caption", 9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
                
        elif b_type == "table":
            lines = text.strip().split("\n")
            if len(lines) >= 2:
                rows_cells = []
                for line in lines:
                    if "---" in line and "|" in line:
                        continue
                    cells = [c.strip() for c in line.split("|")]
                    if cells:
                        rows_cells.append(cells)
                
                if rows_cells:
                    # Springer table captions are ABOVE tables in 9pt
                    add_styled_paragraph(doc, f"Table {element_index}. Table Caption", 9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
                    
                    table = doc.add_table(rows=len(rows_cells), cols=len(rows_cells[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows_cells):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = val
                                for cp in cell.paragraphs:
                                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for cr in cp.runs:
                                        cr.font.name = "Times New Roman"
                                        cr.font.size = Pt(9)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)

    # --- 4. REFERENCES BIBLIOGRAPHY SECTION (9pt Times New Roman) ---
    if references:
        add_styled_paragraph(doc, "References", 12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)
        
        for idx, ref in enumerate(references):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2)
            
            cit_key = ref.get("citation_key") or str(idx + 1)
            r_key = p.add_run(f"{cit_key}. ")
            r_key.font.name = "Times New Roman"
            r_key.font.size = Pt(9)
            r_key.bold = True
            
            raw_ref = ref["raw_reference"]
            raw_ref_clean = re.sub(r"^\[?\d+\]?\s*\.?\s*", "", raw_ref)
            
            r_ref = p.add_run(raw_ref_clean)
            r_ref.font.name = "Times New Roman"
            r_ref.font.size = Pt(9)
